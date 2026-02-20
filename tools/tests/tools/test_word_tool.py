import importlib.util
from pathlib import Path
from unittest.mock import patch

import pytest
from fastmcp import FastMCP
from docx import Document

docx_available = importlib.util.find_spec("docx") is not None
pytestmark = pytest.mark.skipif(not docx_available, reason="python-docx not installed")

if docx_available:
    from aden_tools.tools.word_tool.word_tool import register_tools

TEST_WORKSPACE_ID = "test-workspace"
TEST_AGENT_ID = "test-agent"
TEST_SESSION_ID = "test-session"


@pytest.fixture
def word_tools(mcp: FastMCP, tmp_path: Path):
    with patch("aden_tools.tools.file_system_toolkits.security.WORKSPACES_DIR", str(tmp_path)):
        register_tools(mcp)
        yield {
            "word_generate": mcp._tool_manager._tools["word_generate"].fn,
            "tmp_path": tmp_path,
        }


def test_word_generate_creates_docx(word_tools):
    gen = word_tools["word_generate"]

    doc = {
        "title": "Monthly Analysis",
        "sections": [
            {
                "heading": "Executive Summary",
                "paragraphs": ["Markets were volatile.", "Risk was elevated."],
                "bullets": ["AAPL outperformed", "MSFT stable"],
            },
            {
                "heading": "Key Metrics",
                "paragraphs": [],
                "bullets": [],
                "table": {
                    "columns": ["Metric", "Value"],
                    "rows": [["Drawdown", "-3.2%"], ["Volatility", "18%"]],
                },
            },
        ],
    }

    res = gen(
        path="out/report.docx",
        doc=doc,
        workspace_id=TEST_WORKSPACE_ID,
        agent_id=TEST_AGENT_ID,
        session_id=TEST_SESSION_ID,
    )





    # d = Document(str(out_abs))
    assert res.get("success") is True, res
    out_abs = res.get("output_path")
    assert out_abs, res
    d = Document(str(out_abs))
    full = "\n".join(p.text for p in d.paragraphs)

    assert "Monthly Analysis" in full
    assert "Executive Summary" in full
    assert "Auto-generated report." in full or "Markets were volatile." in full


    assert res.get("metadata", {}).get("sections") == len(doc["sections"])


    out_abs = (
        Path(word_tools["tmp_path"])
        / TEST_WORKSPACE_ID
        / TEST_AGENT_ID
        / TEST_SESSION_ID
        / "out"
        / "report.docx"
    )
    assert out_abs.exists()
    assert out_abs.stat().st_size > 0


def test_word_generate_rejects_bad_extension(word_tools):
    gen = word_tools["word_generate"]
    res = gen(
        path="out/nope.txt",
        doc={"title": "X", "sections": [{"heading": "A"}]},
        workspace_id=TEST_WORKSPACE_ID,
        agent_id=TEST_AGENT_ID,
        session_id=TEST_SESSION_ID,
    )
    assert res.get("success") is False
    assert res.get("error", {}).get("code") == "INVALID_EXTENSION"


def test_word_generate_append_mode(word_tools):
    gen = word_tools["word_generate"]

    res_create = gen(
        path="out/append_report.docx",
        doc={
            "title": "Append Report",
            "sections": [{"heading": "First Section", "paragraphs": ["one"], "bullets": []}],
        },
        workspace_id=TEST_WORKSPACE_ID,
        agent_id=TEST_AGENT_ID,
        session_id=TEST_SESSION_ID,
        mode="create",
    )
    assert res_create.get("success") is True, res_create

    res_append = gen(
        path="out/append_report.docx",
        doc={
            "title": "Ignored In Append",
            "sections": [{"heading": "Second Section", "paragraphs": ["two"], "bullets": []}],
        },
        workspace_id=TEST_WORKSPACE_ID,
        agent_id=TEST_AGENT_ID,
        session_id=TEST_SESSION_ID,
        mode="append",
    )
    assert res_append.get("success") is True, res_append

    out_abs = (
        Path(word_tools["tmp_path"])
        / TEST_WORKSPACE_ID
        / TEST_AGENT_ID
        / TEST_SESSION_ID
        / "out"
        / "append_report.docx"
    )
    d = Document(str(out_abs))
    full = "\n".join(p.text for p in d.paragraphs)
    assert "First Section" in full
    assert "Second Section" in full

