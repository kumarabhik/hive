"""Word Tool - Generate .docx reports (schema-first)."""
from __future__ import annotations
from aden_tools.tools.office_skills_pack.contracts import ArtifactError, ArtifactResult


import os
from typing import Any, List, Optional

from fastmcp import FastMCP
from pydantic import BaseModel, Field

from ..file_system_toolkits.security import get_secure_path


class TableSpec(BaseModel):
    columns: List[str] = Field(..., min_length=1)
    rows: List[List[Any]] = Field(default_factory=list)


class SectionSpec(BaseModel):
    heading: str = Field(..., min_length=1)
    paragraphs: List[str] = Field(default_factory=list)
    bullets: List[str] = Field(default_factory=list)
    table: Optional[TableSpec] = None
    image_paths: List[str] = Field(
        default_factory=list,
        description="Image paths relative to session sandbox",
    )
    charts: List[str] = Field(
        default_factory=list,
        description="Chart PNG paths relative to session sandbox",
    )


class DocSpec(BaseModel):
    title: str = Field(..., min_length=1)
    sections: List[SectionSpec] = Field(..., min_length=1)


def register_tools(mcp: FastMCP) -> None:
    @mcp.tool()
    def word_generate(
        path: str,
        doc: dict[str, Any],
        workspace_id: str,
        agent_id: str,
        session_id: str,
        strict: bool = True,
        mode: str = "create",
    ) -> dict:
        """
        Generate a Word (.docx) report from a strict schema and save it to the session sandbox.
        """

        if not path.lower().endswith(".docx"):
            return ArtifactResult(
                success=False,
                error=ArtifactError(
                    code="INVALID_EXTENSION",
                    message="path must end with .docx",
                ),
            ).model_dump()

        if mode not in {"create", "append"}:
            return ArtifactResult(
                success=False,
                error=ArtifactError(
                    code="INVALID_SCHEMA",
                    message="mode must be 'create' or 'append'",
                ),
            ).model_dump()



        try:
            spec = DocSpec.model_validate(doc)
        except Exception as e:
            return ArtifactResult(
                success=False,
                error=ArtifactError(
                    code="INVALID_SCHEMA",
                    message="Invalid doc schema",
                    details=str(e),
                ),
            ).model_dump()


        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            return ArtifactResult(
                success=False,
                error=ArtifactError(
                    code="DEP_MISSING",
                    message="python-docx not installed. Install with: pip install -e '.[word]' (from tools/)",
                ),
            ).model_dump()
           
            

        try:
            out_path = get_secure_path(path, workspace_id, agent_id, session_id)
            os.makedirs(os.path.dirname(out_path), exist_ok=True)

            if mode == "append":
                if os.path.exists(out_path):
                    d = Document(out_path)
                elif strict:
                    return ArtifactResult(
                        success=False,
                        error=ArtifactError(
                            code="INVALID_PATH",
                            message="docx to append not found",
                        ),
                    ).model_dump()
                else:
                    d = Document()
            else:
                d = Document()
                d.add_heading(spec.title, level=0)
            
            for sec in spec.sections:
                d.add_heading(sec.heading, level=1)

                for p in sec.paragraphs:
                    d.add_paragraph(p)

                for b in sec.bullets:
                    d.add_paragraph(b, style="List Bullet")

                if sec.table is not None:
                    if strict and any(len(row) > len(sec.table.columns) for row in sec.table.rows):
                        return ArtifactResult(
                            success=False,
                            error=ArtifactError(
                                code="INVALID_SCHEMA",
                                message="Table row has more values than columns",
                            ),
                        ).model_dump()
                    t = d.add_table(rows=1, cols=len(sec.table.columns))
                    hdr = t.rows[0].cells
                    for i, c in enumerate(sec.table.columns):
                        hdr[i].text = str(c)

                    for row in sec.table.rows:
                        r = t.add_row().cells
                        for i in range(min(len(r), len(row))):
                            r[i].text = "" if row[i] is None else str(row[i])

                for rel_path in [*sec.image_paths, *sec.charts]:
                    img_path = get_secure_path(rel_path, workspace_id, agent_id, session_id)
                    if not os.path.exists(img_path):
                        if strict:
                            return ArtifactResult(
                                success=False,
                                error=ArtifactError(
                                    code="INVALID_PATH",
                                    message=f"Image not found: {rel_path}",
                                ),
                            ).model_dump()
                        continue
                    d.add_picture(str(img_path), width=Inches(5.8))

                
            d.save(out_path)


            return ArtifactResult(
                success=True,
                output_path=str(out_path),
                metadata={"sections": len(spec.sections)}

            ).model_dump()
                        

        except Exception as e:
            return ArtifactResult(
                success=False,
                error=ArtifactError(
                    code="INTERNAL_ERROR",
                    message="Failed to generate Word document",
                    details=str(e),
                ),
            ).model_dump()
